from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_resume():
    doc = Document()

    # --- Estilos Globales Simples (ATS Friendly) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- 1. ENCABEZADO ---
    header = doc.add_heading('[TU NOMBRE COMPLETO]', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Datos de contacto
    contact_info = doc.add_paragraph()
    contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info.add_run('[Ciudad, País] | [Teléfono] | [Correo Electrónico]\n')
    contact_info.add_run('[Enlace a LinkedIn] | [Enlace a GitHub/Portafolio]')

    # --- 2. PERFIL PROFESIONAL ---
    doc.add_heading('PERFIL PROFESIONAL', level=1)
    profile_text = (
        "Arquitecto de Software y Desarrollador Senior con más de 20 años de experiencia diseñando, "
        "construyendo y escalando sistemas distribuidos de alta disponibilidad. Especialista en modernización "
        "de sistemas heredados (legacy) y adopción de arquitecturas nativas en la nube. Perfil \"Hands-on\": "
        "mantengo una participación activa en la codificación de componentes críticos (aprox. 30-40% del tiempo) "
        "mientras lidero la estrategia técnica. Historial comprobado liderando equipos técnicos, definiendo "
        "estándares de ingeniería y alineando la tecnología con los objetivos comerciales."
    )
    doc.add_paragraph(profile_text)

    # --- 3. HABILIDADES TÉCNICAS ---
    doc.add_heading('HABILIDADES TÉCNICAS', level=1)
    
    skills = [
        ("Arquitectura de Software:", "Microservicios, EDA, Serverless, DDD, Patrones de integración, Clean Arch."),
        ("Lenguajes de Programación:", "Java (v8-v21), C#/.NET Core, Python, Go, JavaScript/TypeScript (Node.js)."),
        ("Nube e Infraestructura:", "AWS (EC2, Lambda, S3, RDS, EKS), Azure, GCP, Terraform (IaC), Docker, K8s."),
        ("Bases de Datos:", "PostgreSQL, SQL Server, MongoDB, Cassandra, Redis, Elasticsearch."),
        ("DevOps y Herramientas:", "CI/CD (Jenkins, GitHub Actions), Kafka, RabbitMQ, Grafana, Prometheus."),
        ("Habilidades Blandas:", "Liderazgo técnico, Mentoría, Gestión de Stakeholders, Agile (Scrum/SAFe).")
    ]

    for category, items in skills:
        p = doc.add_paragraph(style='List Bullet')
        runner = p.add_run(f"{category} ")
        runner.bold = True
        p.add_run(items)

    # --- 4. EXPERIENCIA LABORAL ---
    doc.add_heading('EXPERIENCIA LABORAL', level=1)

    # Trabajo 1
    p = doc.add_paragraph()
    p.add_run('[Nombre de la Empresa Actual] | [Ciudad/Remoto]').bold = True
    p.add_run('\nArquitecto de Software Principal').bold = True
    p.add_run(' | [Mes, Año] – [Presente]')
    
    achievements_1 = [
        "Dirijo la estrategia de arquitectura para una plataforma SaaS que procesa [X] millones de transacciones diarias.",
        "Diseñé y ejecuté la migración de un monolito heredado de 15 años a microservicios en AWS (ahorro 30% OPEX).",
        "Desarrollé en Go/Java el núcleo del motor de procesamiento de pagos, optimizando latencia de 200ms a 50ms.",
        "Implementé 'Infrastructure as Code' con Terraform, estandarizando entornos de desarrollo y producción.",
        "Mentoría técnica a equipo de 15 desarrolladores senior, realizando revisiones de código y guías de estilo."
    ]
    for ach in achievements_1:
        doc.add_paragraph(ach, style='List Bullet')

    # Trabajo 2
    p = doc.add_paragraph()
    p.add_run('\n[Nombre de la Empresa Anterior] | [Ciudad]').bold = True
    p.add_run('\nArquitecto de Soluciones / Tech Lead').bold = True
    p.add_run(' | [Mes, Año] – [Mes, Año]')

    achievements_2 = [
        "Lideré el diseño técnico de [Proyecto], solución crítica para el sector [Banca/Retail].",
        "Introduje Docker y Kubernetes, mejorando la escalabilidad durante eventos de alto tráfico.",
        "Programé componentes de integración en .NET Core para conectar Mainframe con APIs RESTful.",
        "Reduje la deuda técnica en un 40% mediante SonarQube y TDD."
    ]
    for ach in achievements_2:
        doc.add_paragraph(ach, style='List Bullet')

    # Trabajo 3
    p = doc.add_paragraph()
    p.add_run('\n[Nombre de la Empresa Anterior] | [Ciudad]').bold = True
    p.add_run('\nDesarrollador Senior / Líder de Equipo').bold = True
    p.add_run(' | [Mes, Año] – [Mes, Año]')

    achievements_3 = [
        "Arquitecté y desarrollé una aplicación web full-stack utilizada por +50,000 usuarios.",
        "Optimicé consultas SQL Server mejorando en 500% la generación de reportes.",
        "Lideré la transición del equipo de metodología Cascada a Scrum."
    ]
    for ach in achievements_3:
        doc.add_paragraph(ach, style='List Bullet')

    # Trabajo 4 (Resumido por antigüedad)
    p = doc.add_paragraph()
    p.add_run('\n[Nombre de la Empresa - Etapa Temprana] | [Ciudad]').bold = True
    p.add_run('\nDesarrollador de Software (Java/C++)').bold = True
    p.add_run(' | [Mes, Año] – [Mes, Año]')
    doc.add_paragraph("Desarrollo backend para sistemas ERP y ciclo completo SDLC.", style='List Bullet')

    # --- 5. EDUCACIÓN ---
    doc.add_heading('EDUCACIÓN', level=1)
    p = doc.add_paragraph()
    p.add_run('[Título Universitario, ej. Ingeniería de Sistemas]\n').bold = True
    p.add_run('[Nombre de la Universidad], [Ciudad, País]')

    # --- 6. CERTIFICACIONES ---
    doc.add_heading('CERTIFICACIONES', level=1)
    certs = [
        "AWS Certified Solutions Architect – Professional | [Año]",
        "TOGAF 9 Certified | [Año]",
        "Certified Kubernetes Administrator (CKA) | [Año]",
        "Microsoft Certified: Azure Solutions Architect Expert | [Año]"
    ]
    for cert in certs:
        doc.add_paragraph(cert, style='List Bullet')

    # --- 7. IDIOMAS ---
    doc.add_heading('IDIOMAS', level=1)
    doc.add_paragraph('Español: Nativo', style='List Bullet')
    doc.add_paragraph('Inglés: [Nivel, ej. C1 Avanzado / Profesional]', style='List Bullet')

    # Guardar archivo
    file_name = 'Hoja_De_Vida_Arquitecto.docx'
    doc.save(file_name)
    print(f"Archivo '{file_name}' generado exitosamente.")

if __name__ == "__main__":
    create_resume()